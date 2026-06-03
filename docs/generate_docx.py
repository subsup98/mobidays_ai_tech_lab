"""planning.md → planning.docx (편집 가능한 Word 문서)"""
import re
import sys
import os

sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
MD_PATH  = os.path.join(BASE_DIR, "planning.md")
DOCX_PATH = os.path.join(BASE_DIR, "planning.docx")


# ── 헬퍼 ──────────────────────────────────────────────────
def set_font(run, size=None, bold=None, italic=None,
             mono=False, color=None):
    run.font.name = "Courier New" if mono else "맑은 고딕"
    # 동아시아 폰트도 같이 설정
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    font_name = "Courier New" if mono else "맑은 고딕"
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)

    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_run_inline(para, text, bold=False, mono=False, size=10):
    """인라인 텍스트 run 추가."""
    run = para.add_run(text)
    set_font(run, size=size, bold=bold, mono=mono)
    return run


def add_inline_markup(para, text, size=10):
    """**bold**, `code`, **`code`** 를 파싱해 run으로 추가."""
    pattern = re.compile(r'\*\*`([^`]+)`\*\*|\*\*(.+?)\*\*|`([^`]+)`')
    last = 0
    for m in pattern.finditer(text):
        if text[last:m.start()]:
            add_run_inline(para, text[last:m.start()], size=size)
        if m.group(1):        # **`code`**
            add_run_inline(para, m.group(1), bold=True, mono=True, size=size)
        elif m.group(2):      # **bold**
            add_run_inline(para, m.group(2), bold=True, size=size)
        else:                 # `code`
            add_run_inline(para, m.group(3), mono=True, size=size - 0.5)
        last = m.end()
    if text[last:]:
        add_run_inline(para, text[last:], size=size)


def shade_paragraph(para, hex_color="F4F4F4"):
    """문단 배경색 설정 (코드 블록용)."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def set_para_spacing(para, before=0, after=4):
    fmt = para.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after  = Pt(after)


def add_bottom_border(para, color="AAAAAA", size=4):
    """단락 아래 선 (h1 h2 구분선 대용)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── 마크다운 파서 ──────────────────────────────────────────
def parse(doc, md_text):
    lines = md_text.splitlines()
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # H1
        if line.startswith("# "):
            p = doc.add_heading(level=1)
            p.clear()
            add_inline_markup(p, line[2:], size=16)
            add_bottom_border(p, "888888", 6)
            set_para_spacing(p, before=18, after=6)
            i += 1

        # H2
        elif line.startswith("## "):
            p = doc.add_heading(level=2)
            p.clear()
            add_inline_markup(p, line[3:], size=13)
            add_bottom_border(p, "BBBBBB", 4)
            set_para_spacing(p, before=14, after=5)
            i += 1

        # H3
        elif line.startswith("### "):
            p = doc.add_heading(level=3)
            p.clear()
            add_inline_markup(p, line[4:], size=11)
            set_para_spacing(p, before=10, after=4)
            i += 1

        # HR
        elif stripped == "---":
            p = doc.add_paragraph()
            add_bottom_border(p, "CCCCCC", 4)
            set_para_spacing(p, before=6, after=6)
            i += 1

        # Blockquote
        elif line.startswith("> "):
            p = doc.add_paragraph(style="Quote")
            add_inline_markup(p, line[2:], size=10)
            set_para_spacing(p, before=2, after=4)
            i += 1

        # Fenced code block
        elif stripped.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                p = doc.add_paragraph()
                run = p.add_run(lines[i])
                set_font(run, size=9, mono=True)
                shade_paragraph(p)
                fmt = p.paragraph_format
                fmt.space_before = Pt(0)
                fmt.space_after  = Pt(0)
                fmt.left_indent  = Cm(0.5)
                i += 1
            i += 1  # closing ```
            # 코드 블록 후 여백
            sp = doc.add_paragraph()
            set_para_spacing(sp, before=0, after=4)

        # Table
        elif stripped.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl_lines.append(lines[i].strip())
                i += 1
            # 구분선 제외
            rows = [r for r in tbl_lines
                    if not re.match(r"^\|[\s\-|]+\|$", r)]
            if not rows:
                continue
            col_n = max(len(r.strip().strip("|").split("|")) for r in rows)
            table = doc.add_table(rows=len(rows), cols=col_n)
            table.style = "Table Grid"
            for ri, row_txt in enumerate(rows):
                cells = [c.strip() for c in row_txt.strip("|").split("|")]
                while len(cells) < col_n:
                    cells.append("")
                for ci, cell_txt in enumerate(cells):
                    cell = table.cell(ri, ci)
                    cell.paragraphs[0].clear()
                    p = cell.paragraphs[0]
                    add_inline_markup(p, cell_txt, size=9)
                    if ri == 0:
                        for run in p.runs:
                            run.font.bold = True
                        # 헤더 행 배경
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement("w:shd")
                        shd.set(qn("w:val"), "clear")
                        shd.set(qn("w:color"), "auto")
                        shd.set(qn("w:fill"), "EEEEEE")
                        tcPr.append(shd)
            doc.add_paragraph()  # 테이블 후 여백

        # Bullet list
        elif re.match(r"^[-*] ", line):
            while i < len(lines) and re.match(r"^[-*] ", lines[i]):
                p = doc.add_paragraph(style="List Bullet")
                add_inline_markup(p, lines[i][2:], size=10)
                set_para_spacing(p, before=1, after=1)
                i += 1

        # Numbered list
        elif re.match(r"^\d+\. ", line):
            while i < len(lines) and re.match(r"^\d+\. ", lines[i]):
                m = re.match(r"^\d+\. (.+)", lines[i])
                p = doc.add_paragraph(style="List Number")
                add_inline_markup(p, m.group(1), size=10)
                set_para_spacing(p, before=1, after=1)
                i += 1

        # Empty line
        elif not stripped:
            i += 1

        # Normal paragraph
        else:
            parts = []
            while i < len(lines):
                l = lines[i]
                s = l.strip()
                if (not s or l.startswith("#") or s == "---"
                        or s.startswith("|") or s.startswith("```")
                        or l.startswith("> ")
                        or re.match(r"^[-*] ", l)
                        or re.match(r"^\d+\. ", l)):
                    break
                parts.append(l.strip())
                i += 1
            if parts:
                p = doc.add_paragraph()
                add_inline_markup(p, " ".join(parts), size=10)
                set_para_spacing(p, before=2, after=4)


def main():
    with open(MD_PATH, encoding="utf-8") as f:
        md_text = f.read()

    doc = Document()

    # 페이지 여백 설정
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # 기본 폰트 설정
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    parse(doc, md_text)
    doc.save(DOCX_PATH)
    print(f"DOCX 생성 완료: {DOCX_PATH}")


if __name__ == "__main__":
    main()
